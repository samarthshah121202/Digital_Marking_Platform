import os
import pandas as pd
from django.conf import settings
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import fitz  # PyMuPDF
import re
import tabula
import logging
from docx import Document 
from pypdf import PdfReader
from openpyxl import load_workbook
from operator import attrgetter
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404
from django.urls import reverse, path 
from django.http import Http404, FileResponse
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
import os
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, PatternFill, Border, Side




from main.models import Feedback, Module, Question, Section
logger = logging.getLogger(__name__)

#def read_group_number(pdf_path):
    


def extract_student_info_from_pdf(pdf_path, is_group=False):
    # Use tabula to extract tables from the first page of the PDF
    reader = PdfReader(pdf_path)
    page = reader.pages[0]
    text = page.extract_text()
    txt_arr = text.split(" ")
    for i in range(len(txt_arr)):
        if txt_arr[i] == "\nGroup":
            group_num = txt_arr[i + 1]
    
    
    
    tables = tabula.read_pdf(pdf_path, pages=1, multiple_tables=True)
    #("group number: ", group_num)

    student_info = []

    if tables:
        # Assuming the first table is the one that contains the student data
        # You may need to inspect the tables to ensure this is correct
        table = tables[0]
      #  print(table)
        #logger.info(table)

        first_name = None
        last_name = None
        student_number = None
       # group_number = None 

        #logger.info(f"is_group: {is_group}")    

        # Assuming the table has columns for first name, last name, and student number
        if is_group:
            for index, row in table.iterrows():
               # print(index, row)
                
 
                
                logger.error(f"row: {row}")  

                last_name = row[0]
                first_name = row[1]
                student_number = row[2]
                #group_number = row[3]
                group_number = group_num
                #logger.info(f"first_name: {first_name}, last_name: {last_name}, student_number: {student_number}, group_number: {group_number}")
                student_info.append({
                    "first_name": first_name,
                    "last_name": last_name,
                    "student_number": student_number,
                    "group_number": group_number
                })
            #    print(student_info)
        else:
            for index, row in table.iterrows():
               # print(index, row)
                # Extract the first name, last name, and student number
                if row[0] == 'First Name':
                    first_name = row[1]
                elif row[0] == 'Last Name':
                    last_name = row[1]
                elif row[0] == 'Student ID':
                    student_number = row[1]
                #logger.info(f"first_name: {first_name}, last_name: {last_name}, student_number: {student_number}")

            if first_name and last_name and student_number:
                student_info.append({
                    "first_name": first_name,
                    "last_name": last_name,
                    "student_number": student_number,
                })

    #logger.info(f"RETURNING FROM extract_student_info_from_pdf: {student_info} {type(student_info)}")
    return student_info

def handle_uploaded_file(f, filepath):
    folder_path = os.path.dirname(filepath)
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    with open(filepath, "wb+") as destination:
        for chunk in f.chunks():
            destination.write(chunk)
    
def create_markscheme_objects(assignment, markscheme_data_frame):
    sections = []
    modules = []
    questions = []
    feedbacks = []
    
    current_section = None
    current_module = None
    current_question = None

    for index, row in markscheme_data_frame.iterrows():
        row_values = row.values.tolist()
        
        # Skip empty rows
        if all(pd.isna(value) for value in row_values):
            continue
            
        # Check first column for section/topic/question
        if pd.notna(row_values[0]):
            first_col = str(row_values[0]).strip()
            
            # Create Section
            if "Part" in first_col:
                current_section = Section.objects.create(
                    section_name=first_col,
                    assignment=assignment
                )
                sections.append(current_section)
                current_module = None  # Reset current module
                current_question = None  # Reset current question
                
            # Create Module (Topic)
            elif "Topic" in first_col:
                if current_section:
                    current_module = Module.objects.create(
                        module_name=first_col,
                        section=current_section
                    )
                    modules.append(current_module)
                    current_question = None  # Reset current question
                    
            # Create Question
            elif "Question" in first_col:
                if current_module:
                    current_question = Question.objects.create(
                        question=first_col,
                        module=current_module
                    )
                    questions.append(current_question)
        
        # Create Feedback
        if current_question and pd.notna(row_values[1]):
            try:
                feedback_key = str(row_values[1]).strip()
                feedback_text = str(row_values[2]).strip() if pd.notna(row_values[2]) else ""
                mark = float(row_values[3]) if pd.notna(row_values[3]) else 0.0
                
                feedback = Feedback.objects.create(
                    feedback_key=feedback_key,
                    feedback_text=feedback_text,
                    mark=mark,
                    question=current_question
                )
                feedbacks.append(feedback)
                
            except Exception as e:
                logger.error(f"Error creating feedback for question {current_question.question}: {str(e)}")
                logger.error(f"Row values: {row_values}")

    #logger.info(f"sections:{sections}")
    return sections

def handle_upload_excel_sheet(filePath, save_folder, filename, assignment) -> list[Section]:
    excel_file = pd.ExcelFile(filePath)
    sheet_names = excel_file.sheet_names
    df = pd.read_excel(filePath, sheet_name=sheet_names[0])

    csv_filename = os.path.join(save_folder,"markscheme",f"{filename}.csv")
    markscheme_obj = create_markscheme_objects(assignment, df)

    df.to_csv(csv_filename, index=False, encoding='cp1252')

    return markscheme_obj

def create_feedback_doc(student_infos, sections, assignment_title, student_feedback_doc_path, assignment, student_work, feedback_above_50, feedback_below_50):
    doc = Document()

    # Set the document orientation to landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Modify title to include "- Assignment Feedback" and apply formatting
    title_text = f"{assignment_title} - Assignment Feedback"
    title_paragraph = doc.add_heading(level=1)
    title_run = title_paragraph.add_run(title_text)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
    title_paragraph.alignment = 1  # Center alignment

    # Add feedback summary table (different format for group vs individual assignments)
    if assignment.is_group_assignment:
        # Group assignment: 3-column table with group number
        feedback_table = doc.add_table(rows=1, cols=3)
        feedback_table.style = 'Table Grid'
        feedback_table.cell(0, 0).text = f"Group Number: {student_work.group_number}"
        feedback_table.cell(0, 1).text = "Strengths of the Report"
        feedback_table.cell(0, 2).text = "Areas that Need Improving"

        # Apply colors to relevant columns
        feedback_table.cell(0, 1)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="C6E0B4"/>'.format(nsdecls('w'))))  # Pastel Green
        feedback_table.cell(0, 2)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFF2CC"/>'.format(nsdecls('w'))))  # Pastel Yellow
    else:
        # Individual assignment: 2-column table without group number
        feedback_table = doc.add_table(rows=1, cols=2)
        feedback_table.style = 'Table Grid'
        feedback_table.cell(0, 0).text = "Strengths of the Report"
        feedback_table.cell(0, 1).text = "Areas that Need Improving"

        # Apply colors to relevant columns
        feedback_table.cell(0, 0)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="C6E0B4"/>'.format(nsdecls('w'))))  # Pastel Green
        feedback_table.cell(0, 1)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFF2CC"/>'.format(nsdecls('w'))))  # Pastel Yellow

    # Add some space between the feedback summary table and the student info table
    doc.add_paragraph()

    # Determine the number of columns for the student table
    num_cols = len(student_infos[0]) - (1 if assignment.is_group_assignment else 0)

    # Create and populate the student information table
    student_table = doc.add_table(rows=len(student_infos), cols=num_cols)
    student_table.style = 'Table Grid'

    # Populate the table with student info
    for row_idx, row_data in enumerate(student_infos):
        for col_idx in range(num_cols):
            student_table.cell(row_idx, col_idx).text = str(row_data[col_idx])

    # Add a blank paragraph to separate the student information table from the feedback table
    doc.add_paragraph()

    # Create feedback table with a single column
    feedback_table = doc.add_table(rows=0, cols=1)
    feedback_table.style = 'Table Grid'

    # Populate the feedback table
    for section in sections:
        section_name = section["section"].section_name

        # Add row for section name
        section_row = feedback_table.add_row().cells
        section_row[0].text = section_name

        for module in section["modules"]:
            module_name = module["module"].module_name

            # Add row for module name (topic)
            module_row = feedback_table.add_row().cells
            module_row[0].text = module_name

        
            feedback_above_list = []
            feedback_below_list = []

        # Process each question in the module
            for question in module["questions"]:
                feedback_text = question["feedback_text"]
                custom_feedback = question["custom_feedback"] or ""



                full_feedback = f"{feedback_text} {custom_feedback}".strip()

                # Categorize feedback
                if feedback_text in feedback_above_50:
                    feedback_above_list.append(full_feedback)
                elif feedback_text in feedback_below_50:
                    feedback_below_list.append(full_feedback)

        # Append feedback rows for the module
            if feedback_above_list:
                feedback_above_row = feedback_table.add_row().cells
                feedback_above_row[0].text = "\n".join(feedback_above_list)
                feedback_above_row[0]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="C6E0B4"/>'.format(nsdecls('w'))))


            if feedback_below_list:
                feedback_below_row = feedback_table.add_row().cells
                feedback_below_row[0].text = "\n".join(feedback_below_list)
                feedback_below_row[0]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFF2CC"/>'.format(nsdecls('w'))))





    for row_index, row in enumerate(feedback_table.rows):
        for cell in row.cells:
            if "Part" in cell.text:
                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="AEC6CF"/>'.format(nsdecls('w'))))
                # Format text: Center, Bold, and Larger Font
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center text
                
                run = paragraph.runs[0]
                run.bold = True  # Make text bold
                run.font.size = Pt(12)  # Increase font size
                break  # Stop checking this row after the first match

    for row_index, row in enumerate(feedback_table.rows):
        for cell in row.cells:
            if "Topic" in cell.text:            
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.bold = True
                break

    # Define full path and ensure directory exists
    first_name = student_infos[1][0]
    last_name = student_infos[1][1]
    student_id = student_infos[1][2]

    if assignment.is_group_assignment:
        filename = f"Group_{student_work.group_number}_Student_Feedback.docx"
    else:
        filename = f"{first_name}_{last_name}_{student_id}_Student_Feedback.docx"

    full_path = os.path.join(student_feedback_doc_path, filename)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    doc.save(full_path)

    return full_path

def add_to_feedback_sheet(workbook, id_table, group_table=None):

    def add_table(table, sheet_name, add_line=False):
        print("THIS IS THE TABLE", table)
        sheet = workbook[sheet_name]
        if add_line is True:
            table.append([" "])
        for row_data in table:
            new_row = sheet.max_row + 1
            for col_num, cell in enumerate(row_data, start=1):
                #logger.info(f"row={new_row} col={col_num} cell_data={cell} max_row={sheet.max_row}")
                sheet.cell(row=new_row, column=col_num).value = cell

        for col in range(1, sheet.max_column + 1):  # Starting from column 2
            col_letter = get_column_letter(col)
            sheet.column_dimensions[col_letter].width = 20  # Adjust width as needed

    # Adjust row heights - Set height for all rows
        for row in range(1, sheet.max_row + 1):
            sheet.row_dimensions[row].height = 20  #


        thin_border = Border(
        top=Side(style='thin'),
        bottom=Side(style='thin'),
        left=Side(style='thin'),
        right=Side(style='thin')
    )
    
        for row in range(1, sheet.max_row + 1):
            for col in range(1, sheet.max_column + 1):
                cell = sheet.cell(row=row, column=col)
                if cell.value is not None:  # Only apply borders to cells with data
                    cell.border = thin_border        

    add_table(id_table, "Id List")
    
    # Only add group_table if it is provided (not None)
    if group_table is not None:
        add_table(group_table, "Group List", add_line=True)
    
    #add_to_marks_breakdown()
    return 


def question_mark_excel(workbook, processed_questions, processed_modules, processed_sections, group_number, total_marks):
   # print("QUrkESTION MARK EXCEL CALED")
    
    marks_breakdown_sheet = workbook["Marks Breakdown"]

    # Adjust column widths
    # Set a larger width for column 1 (e.g., 30)
    marks_breakdown_sheet.column_dimensions["A"].width = 50  # Column 1 is "A"

    # Set width for other columns (e.g., 15 for columns 2, 3, etc.)
    for col in range(2, marks_breakdown_sheet.max_column + 1):  # Starting from column 2
        col_letter = get_column_letter(col)
        marks_breakdown_sheet.column_dimensions[col_letter].width = 15  # Adjust width as needed

    # Adjust row heights - Set height for all rows
    for row in range(1, marks_breakdown_sheet.max_row + 1):
        marks_breakdown_sheet.row_dimensions[row].height = 20  # Adjust height as neede

    bold_font = Font(bold=True)
    
    # Make all cells in row 1 bold
    for col in range(1, marks_breakdown_sheet.max_column + 1):
        marks_breakdown_sheet.cell(row=1, column=col).font = bold_font
    
    # Make all cells in column 1 bold
    for row in range(1, marks_breakdown_sheet.max_row + 1):
        marks_breakdown_sheet.cell(row=row, column=1).font = bold_font

    marks = []  # {{ edit_3 }}
    
    for section in processed_sections:
        marks.append({section["total"]})  
        for module in section["modules"]:  
            marks.append({module["total"]})  

            for question in module["questions"]:  
                marks.append({question["mark"]}) 
    
    group_col = 1
    row_num = 1
    for col in marks_breakdown_sheet[row_num]:

        if col.value == "Group " + str(group_number):

            break
        else:
            group_col += 1

    list_of_list_marks = [[list(d)[0]] for d in marks]
    list_of_marks = [item[0] for item in list_of_list_marks]
    print("Marks:", list_of_marks) 

    for row_num, item in enumerate(list_of_marks, start=2):
       marks_breakdown_sheet.cell(row=row_num, column=group_col, value=item)

    pastel_yellow_fill = PatternFill(start_color="FFFFE0", end_color="FFFFE0", fill_type="solid")  # Pastel yellow color
    pastel_green_fill = PatternFill(start_color="ACE1AF", end_color="ACE1AF", fill_type="solid")  # Pastel green color
    light_blue_fill = PatternFill(start_color="F0F8FF", end_color="F0F8FF", fill_type="solid")  # Light Blue
    light_grey_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")  # Light Grey


    for row in range(1, marks_breakdown_sheet.max_row + 1):
        cell_value = marks_breakdown_sheet.cell(row=row, column=1).value
        if cell_value and "part" in str(cell_value).lower():  # Case-insensitive check for "part"
            # Fill the entire row with pastel yellow color
            for col in range(1, marks_breakdown_sheet.max_column + 1):  # Loop through all columns in the row
                marks_breakdown_sheet.cell(row=row, column=col).fill = pastel_yellow_fill

        elif "topic" in str(cell_value).lower():
                print(f"Found 'topic' in row {row}, column A: {cell_value}")  # Print the row and value
                # Fill the entire row with pastel green color
                for col in range(1, marks_breakdown_sheet.max_column + 1):
                    marks_breakdown_sheet.cell(row=row, column=col).fill = pastel_green_fill

    color_toggle = True  # This will toggle between True (Light Blue) and False (Light Grey)
    for row in range(1, marks_breakdown_sheet.max_row + 1):
        cell_value = marks_breakdown_sheet.cell(row=row, column=1).value
        if cell_value and "question" in str(cell_value).lower():
            #rint(f"Found 'question' in row {row}, column 1: {cell_value}")

            # Decide which fill color to use (Light Blue or Light Grey)
            if color_toggle:
                fill_color = light_blue_fill
                color_toggle = False  # Toggle to next color (Light Grey)
            else:
                fill_color = light_grey_fill
                color_toggle = True  # Toggle to next color (Light Blue)

            # Apply the chosen fill color and border to the cell
            for col in range(1, marks_breakdown_sheet.max_column + 1):
                cell = marks_breakdown_sheet.cell(row=row, column=col)
                cell.fill = fill_color

    group_col = None
    for idx, col in enumerate(marks_breakdown_sheet[1], start=1):  # Start at column index 1
        if col.value == "Group " + str(group_number):
            group_col = idx
            break

    total_row = None
    for row in range(1, marks_breakdown_sheet.max_row + 1):
        if marks_breakdown_sheet.cell(row=row, column=1).value == "Total":
            total_row = row
            break

    marks_breakdown_sheet.cell(row=total_row, column=group_col, value=total_marks)
    

    thin_border = Border(
        top=Side(style='thin'),
        bottom=Side(style='thin'),
        left=Side(style='thin'),
        right=Side(style='thin')
    )
    
    for row in range(1, marks_breakdown_sheet.max_row + 1):
        for col in range(1, marks_breakdown_sheet.max_column + 1):
            cell = marks_breakdown_sheet.cell(row=row, column=col)
            if cell.value is not None:  # Only apply borders to cells with data
                cell.border = thin_border
                
    total_marks = 0


def question_mark_excel_student(workbook, processed_questions, processed_modules, processed_sections, student_number, total_marks):
   # print("QUrkESTION MARK EXCEL CALED")
    marks_breakdown_sheet = workbook["Marks Breakdown"]
     # Adjust column widths
    # Set a larger width for column 1 (e.g., 30)
    marks_breakdown_sheet.column_dimensions["A"].width = 50  # Column 1 is "A"

    # Set width for other columns (e.g., 15 for columns 2, 3, etc.)
    for col in range(2, marks_breakdown_sheet.max_column + 1):  # Starting from column 2
        col_letter = get_column_letter(col)
        marks_breakdown_sheet.column_dimensions[col_letter].width = 15  # Adjust width as needed

    # Adjust row heights - Set height for all rows
    for row in range(1, marks_breakdown_sheet.max_row + 1):
        marks_breakdown_sheet.row_dimensions[row].height = 20  # Adjust height as neede

    bold_font = Font(bold=True)
    
    # Make all cells in row 1 bold
    for col in range(1, marks_breakdown_sheet.max_column + 1):
        marks_breakdown_sheet.cell(row=1, column=col).font = bold_font
    
    # Make all cells in column 1 bold
    for row in range(1, marks_breakdown_sheet.max_row + 1):
        marks_breakdown_sheet.cell(row=row, column=1).font = bold_font

    marks = []  # {{ edit_3 }}
   # print("student number is: ", student_number)
    
    for section in processed_sections:
        marks.append({section["total"]})  
        for module in section["modules"]:  
            marks.append({module["total"]})  

            for question in module["questions"]:  
                marks.append({question["mark"]}) 
    
    student_col = 1
    row_num = 1
    for col in marks_breakdown_sheet[row_num]:
        print(col.value)

        if col.value == student_number:

            break
        else:
            student_col += 1

    list_of_list_marks = [[list(d)[0]] for d in marks]
    list_of_marks = [item[0] for item in list_of_list_marks]
    print("Marks:", list_of_marks) 

    for row_num, item in enumerate(list_of_marks, start=2):
       marks_breakdown_sheet.cell(row=row_num, column=student_col, value=item)
    pastel_yellow_fill = PatternFill(start_color="FFD700", end_color="FFD700", fill_type="solid")  # Pastel yellow color
    pastel_green_fill = PatternFill(start_color="66CC66", end_color="66CC66", fill_type="solid")  # Pastel green color
    light_blue_fill = PatternFill(start_color="ADD8E6", end_color="ADD8E6", fill_type="solid")  # Light Blue
    light_grey_fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")  # Light Grey


    for row in range(1, marks_breakdown_sheet.max_row + 1):
        cell_value = marks_breakdown_sheet.cell(row=row, column=1).value
        if cell_value and "part" in str(cell_value).lower():  # Case-insensitive check for "part"
            # Fill the entire row with pastel yellow color
            for col in range(1, marks_breakdown_sheet.max_column + 1):  # Loop through all columns in the row
                marks_breakdown_sheet.cell(row=row, column=col).fill = pastel_yellow_fill

        elif "topic" in str(cell_value).lower():
                print(f"Found 'topic' in row {row}, column A: {cell_value}")  # Print the row and value
                # Fill the entire row with pastel green color
                for col in range(1, marks_breakdown_sheet.max_column + 1):
                    marks_breakdown_sheet.cell(row=row, column=col).fill = pastel_green_fill

    color_toggle = True  # This will toggle between True (Light Blue) and False (Light Grey)
    for row in range(1, marks_breakdown_sheet.max_row + 1):
        cell_value = marks_breakdown_sheet.cell(row=row, column=1).value
        if cell_value and "question" in str(cell_value).lower():
            #rint(f"Found 'question' in row {row}, column 1: {cell_value}")

            # Decide which fill color to use (Light Blue or Light Grey)
            if color_toggle:
                fill_color = light_blue_fill
                color_toggle = False  # Toggle to next color (Light Grey)
            else:
                fill_color = light_grey_fill
                color_toggle = True  # Toggle to next color (Light Blue)

            # Apply the chosen fill color and border to the cell
            for col in range(1, marks_breakdown_sheet.max_column + 1):
                cell = marks_breakdown_sheet.cell(row=row, column=col)
                cell.fill = fill_color

    student_col = None
    for idx, col in enumerate(marks_breakdown_sheet[1], start=1):  # Start at column index 1
        if col.value == str(student_number):
            student_col = idx
            break

    total_row = None
    for row in range(1, marks_breakdown_sheet.max_row + 1):
        if marks_breakdown_sheet.cell(row=row, column=1).value == "Total":
            total_row = row
            break

    marks_breakdown_sheet.cell(row=total_row, column=student_col, value=str(total_marks))
    
    

    thin_border = Border(
        top=Side(style='thin'),
        bottom=Side(style='thin'),
        left=Side(style='thin'),
        right=Side(style='thin')
    )
    
    for row in range(1, marks_breakdown_sheet.max_row + 1):
        for col in range(1, marks_breakdown_sheet.max_column + 1):
            cell = marks_breakdown_sheet.cell(row=row, column=col)
            if cell.value is not None:  # Only apply borders to cells with data
                cell.border = thin_border
                
    total_marks = 0
         
       
